# -*- coding: utf-8 -*-
"""Genera Memo_Seguridad_Cifrado_PII_CRM_ASICXXI.docx"""
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x00, 0x49, 0x75)
CYAN = RGBColor(0x00, 0x73, 0xA8)
GREEN = RGBColor(0x16, 0x7A, 0x3C)
AMBER = RGBColor(0xB4, 0x5B, 0x08)

d = docx.Document()
normal = d.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
for lvl, sz in [('Heading 1', 15), ('Heading 2', 12.5), ('Heading 3', 11)]:
    st = d.styles[lvl]
    st.font.size = Pt(sz); st.font.color.rgb = NAVY; st.font.name = 'Calibri'

def h1(t): d.add_heading(t, level=1)
def h2(t): d.add_heading(t, level=2)
def p(t='', bold=False, italic=False, color=None):
    par = d.add_paragraph(); r = par.add_run(t)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return par
def bullet(t):
    par = d.add_paragraph(style='List Bullet'); par.add_run(t); return par
def num(t):
    par = d.add_paragraph(style='List Number'); par.add_run(t); return par

def table(headers, rows, widths=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    d.add_paragraph()
    return t

# ===== PORTADA =====
title = d.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Memo técnico de seguridad'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Cifrado de datos personales en el CRM ASICXXI'); r.font.size = Pt(14); r.font.color.rgb = CYAN
sub2 = d.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run('Agosto 2026 · Respuesta al requisito ENS planteado por IT').italic = True
d.add_paragraph()

h1('1. Motivo')
p('IT planteó un requisito correcto y necesario en el marco de aplicación del Esquema Nacional de '
  'Seguridad (ENS): los datos de carácter personal gestionados por el CRM deben estar protegidos de forma '
  'que, en caso de acceso no autorizado a la base de datos (credenciales filtradas, inyección SQL, backup '
  'sustraído, o compromiso de la cuenta de la plataforma de despliegue), la información no sea legible.')
p('Este documento resume la solución implementada, su alcance, y lo que queda como mejora recomendada a futuro.')

h1('2. Resumen ejecutivo')
p('Se ha implementado cifrado de datos personales a nivel de columna en la base de datos (no solo cifrado '
  'de disco). Está desplegado y validado en los entornos de desarrollo y producción. Todos los datos '
  'personales existentes, además de los nuevos, están cifrados. Los backups anteriores a la migración, que '
  'contenían datos en claro, han sido eliminados de la infraestructura de Railway; se conservó una copia '
  'de seguridad previa a la migración, cifrada como fichero independiente, fuera de la plataforma de '
  'despliegue, como red de seguridad.')

h1('3. Por qué el cifrado de disco no era suficiente')
p('La infraestructura de despliegue (Railway) cifra los discos a nivel de proveedor. Esa protección cubre '
  'exclusivamente el robo físico del disco, un escenario de baja probabilidad. No protege frente al acceso '
  'lógico a los datos, que es el riesgo real:')
table(['Escenario de acceso lógico', '¿Cifrado de disco protege?', '¿Cifrado en columna protege?'], [
    ['Filtración de credenciales de la base de datos', 'No', 'Sí'],
    ['Inyección SQL a través de la aplicación', 'No', 'Sí'],
    ['Compromiso de la cuenta/proyecto de la plataforma de despliegue', 'No', 'Sí'],
    ['Backup sustraído o mal gestionado', 'No', 'Sí'],
], widths=[3.6, 1.5, 1.5])
p('En los cuatro casos, si las columnas con datos personales están cifradas, lo que obtiene quien accede es '
  'una cadena ilegible. La clave de descifrado no reside en la base de datos ni en los backups.')

h1('4. Solución implementada')
h2('4.1 Cifrado en columna')
bullet('Algoritmo: Fernet (AES-128 en modo CBC + HMAC-SHA256 para integridad), mediante la librería estándar "cryptography".')
bullet('Implementación transparente: un TypeDecorator de SQLAlchemy cifra el valor al escribir y lo descifra al leer. El resto del código de la aplicación no necesita cambios.')
bullet('Migración gradual segura: durante el despliegue, si un valor almacenado no es un token cifrado válido (dato heredado aún no migrado), se interpreta como texto en claro sin error. Esto permitió desplegar sin tiempo de inactividad y migrar los datos existentes de forma controlada, con verificación en cada paso.')

h2('4.2 Campos cifrados')
table(['Tabla', 'Campos'], [
    ['Cuentas (empresas cliente)', 'Email, teléfono, dirección, CIF/NIF'],
    ['Contactos', 'Nombre, apellidos'],
    ['Canales de contacto', 'Valor del canal (email o teléfono)'],
    ['Usuarios', 'Firma de email (contiene teléfono/email personal)'],
], widths=[2.6, 4.2])

h2('4.3 Campos NO cifrados — justificación')
p('La razón social de la cuenta (nombre de empresa) y el nombre de la oportunidad comercial se mantienen '
  'en claro de forma deliberada:')
bullet('La razón social es un dato de naturaleza pública (consta en BORME y registros mercantiles); su nivel de sensibilidad es bajo.')
bullet('Es el campo sobre el que se realizan las búsquedas y la ordenación en la aplicación (buscador de clientes, tablero comercial). Cifrarlo impediría la búsqueda por texto parcial sin una reingeniería mayor.')
p('El resto de datos que sí identifican o permiten contactar a una persona física están cifrados en su totalidad.')

h2('4.4 Gestión de claves')
bullet('Dos claves independientes por entorno: una para el cifrado (Fernet) y otra para índices de búsqueda ciega (HMAC-SHA256), que permiten localizar registros por coincidencia exacta sin exponer el valor.')
bullet('Desarrollo y producción usan claves distintas y no compartidas.')
bullet('Las claves se almacenan como variables de entorno en la plataforma de despliegue, fuera del código fuente y fuera del control de versiones.')
bullet('Las claves no residen en la base de datos ni en los backups: un volcado de la base de datos, por sí solo, no permite el descifrado.')
p('Nota de transparencia: el nivel de aislamiento actual (variable de entorno en la misma plataforma que la '
  'base de datos) protege frente a los cuatro escenarios de la sección 3. No protege frente al escenario '
  'extremo de que la cuenta completa de la plataforma de despliegue quede comprometida, ya que en ese caso '
  'el atacante también accedería a las variables de entorno. Para ese nivel de aislamiento se recomienda la '
  'medida descrita en la sección 6 (KMS externo).', italic=True)

h2('4.5 Conexión a base de datos')
p('Se ha incorporado soporte opcional para exigir SSL en la conexión a PostgreSQL, configurable por variable '
  'de entorno, como capa adicional de cifrado en tránsito.')

h1('5. Migración de datos y limpieza de backups')
num('Se validó el proceso completo primero en el entorno de desarrollo: migración de datos históricos, verificación de que la aplicación seguía funcionando con normalidad (búsquedas, edición de fichas, importación) y confirmación de que los valores quedaban cifrados en la base de datos.')
num('Se tomó una copia de seguridad completa de producción previa a la migración, cifrada como fichero independiente (AES-256) y almacenada fuera de la plataforma de despliegue, con verificación de integridad.')
num('Se ejecutó la migración en producción en modo simulación (recuento sin aplicar cambios), validando las cifras antes de aplicar.')
num('Se aplicó la migración en producción. Verificación posterior: cero registros pendientes de cifrar.')
num('Se eliminaron los backups automáticos de producción anteriores a la migración, que contenían los datos personales en claro. Se conservó únicamente el backup automático más reciente dentro de la política de retención habitual, que expirará de forma natural.')
p('A partir de este punto, cualquier backup automático nuevo (diario, semanal o mensual) captura la base de '
  'datos con los datos personales ya cifrados, sin intervención adicional.')

h1('6. Salvaguardas adicionales recomendadas (no implementadas aún)')
p('El cifrado en columna resuelve el requisito planteado. Para reforzar la postura de seguridad en línea con '
  'un esquema ENS más exigente, se recomienda valorar:')
bullet('Autenticación en dos factores (2FA) para las cuentas con rol administrador.')
bullet('Gestión de claves mediante un servicio externo dedicado (KMS: AWS KMS, GCP KMS, HashiCorp Vault, Doppler…), de modo que la clave de cifrado no resida en la misma plataforma que la base de datos.')
bullet('Auditoría de que no se registran datos personales en los logs de la aplicación.')
bullet('Política formal de retención y borrado de datos personales.')
bullet('Aplicar el mismo patrón de cifrado en columna a la herramienta de Onboarding y a cualquier aplicación futura que gestione datos de clientes.')
bullet('Rotación periódica documentada de las claves de cifrado.')

h1('7. Alcance de este documento')
p('Este memo cubre la dimensión técnica del requisito. El Esquema Nacional de Seguridad exige además un '
  'marco formal — categorización del sistema (nivel BAJO/MEDIO/ALTO), análisis de riesgos y Declaración de '
  'Aplicabilidad — que corresponde definir con el responsable de seguridad de la organización.')

h1('8. Conclusión')
p('Los datos personales del CRM ASICXXI (nombre, apellidos, email, teléfono, dirección y CIF/NIF de clientes '
  'y contactos) están cifrados a nivel de aplicación, tanto los históricos como los nuevos, en desarrollo y '
  'en producción. Un volcado de la base de datos o un backup sustraído resulta ilegible sin la clave, que no '
  'reside en ninguno de los dos.')

d.add_paragraph()
end = d.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.add_run('— Fin del documento —').italic = True

out = 'Memo_Seguridad_Cifrado_PII_CRM_ASICXXI.docx'
d.save(out)
print('GENERADO:', out)
