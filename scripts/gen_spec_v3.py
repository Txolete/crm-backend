# -*- coding: utf-8 -*-
"""Genera CRM_ASICXXI_Spec_ClaudeCode_v3.0.docx consolidando V2.1 + sesion + N1 + N5v2."""
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x00, 0x49, 0x75)
CYAN = RGBColor(0x00, 0x73, 0xA8)

d = docx.Document()

# Estilos base
normal = d.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

for lvl, sz in [('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11.5)]:
    st = d.styles[lvl]
    st.font.size = Pt(sz)
    st.font.color.rgb = NAVY
    st.font.name = 'Calibri'

def h1(t): d.add_heading(t, level=1)
def h2(t): d.add_heading(t, level=2)
def h3(t): d.add_heading(t, level=3)
def p(t='', bold=False, italic=False):
    par = d.add_paragraph()
    r = par.add_run(t); r.bold = bold; r.italic = italic
    return par
def bullet(t):
    par = d.add_paragraph(style='List Bullet'); par.add_run(t); return par
def num(t):
    par = d.add_paragraph(style='List Number'); par.add_run(t); return par

def table(headers, rows, widths=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    d.add_paragraph()
    return t

# ===== PORTADA =====
title = d.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CRM ASICXXI'); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Especificación Técnica Consolidada — v3.0'); r.font.size = Pt(15); r.font.color.rgb = CYAN
sub2 = d.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run('Junio 2026 · Estado real + Roadmap + N1 Onboarding + N5v2 Business Plan + N2 Comunicaciones').italic = True
d.add_paragraph()
p('La v3.0 sustituye a la v2.1. Integra: (a) el estado real tras los desarrollos de junio 2026 '
  '(feedback in-app, material comercial, plantillas de email, hub de comunicaciones), (b) la spec de '
  'integración del Gestor de Onboarding (N1), y (c) la spec del Business Plan Tool integrado (N5 v2). '
  'Es la única fuente de verdad para los próximos sprints. Cada sprint pendiente puede lanzarse en una '
  'sesión limpia de Claude Code.', italic=True)

d.add_page_break()

# ===== 1. INSTRUCCIONES GENERALES =====
h1('1. Instrucciones generales para Claude Code')
p('Reglas absolutas. Toda sesión debe respetar este bloque antes de empezar cualquier sprint.')
h3('Reglas de trabajo')
for t in [
    'Trabaja siempre sobre la rama develop. Nunca toques main. Ejecuta git branch antes de cada tarea.',
    'Lee cada fichero afectado COMPLETO antes de modificarlo. No asumas su contenido.',
    'Mantén JS vanilla + Jinja2. No introduzcas React, Vue ni ningún framework de frontend.',
    'No dupliques lógica. Las funciones compartidas viven en app/utils/.',
    'Fechas: usa SIEMPRE get_utc_now() en asignaciones a modelos. Columnas con Column(UTCDateTime()) (TypeDecorator en app/database.py).',
    'Single-commit + db.refresh() antes de construir el Response Pydantic (expire_on_commit=True).',
]:
    bullet(t)
h3('Decisiones técnicas (no revertir)')
for t in [
    'IDs: todas las entidades del CRM usan VARCHAR/UUID (string), NO enteros. Cualquier FK nueva (onboardings, business_plans, etc.) debe ser String, no INT.',
    'calculate_probability: función canónica en app/utils/opportunity.py — no duplicar en rutas.',
    'Roles: admin, sales, commercial, viewer. commercial filtra por owner_user_id == current_user.id en kanban, opportunities, tasks, dashboard, accounts (y en cualquier módulo nuevo con datos por owner).',
    'Tablas Boolean en PostgreSQL: convertir bool→int al asignar is_active/is_terminal (PostgreSQL estricto).',
    'Plantillas de correo: existe UNA tabla reutilizable, email_templates (modelo app/models/email_template.py). Cualquier nueva necesidad de plantillas de correo (Business Plan, etc.) reutiliza esta tabla. NO crear cfg_email_templates ni tablas paralelas.',
    'IA: OpenAI vía Responses API. Modelo general gpt-5.4; el hub de comunicaciones usa gpt-5.5 (config comunicaciones_ai_model). gpt-5.5 solo admite temperature=1 (no enviar temperature).',
]:
    bullet(t)
h3('Convención de commits')
for t in ['fix(sprint-N): bug corregido', 'feat(sprint-N): nueva funcionalidad', 'refactor(sprint-N): deuda técnica', 'chore(sprint-N): infra, migraciones, config']:
    bullet(t)
h3('Proceso de deploy a PRO')
for i, t in enumerate([
    'Backup BD PRO.',
    'Merge develop → main (GitHub Desktop, merge commit).',
    'Railway despliega y ejecuta alembic upgrade head como pre-deploy.',
    'Verificar que el deploy pasa a Active sin errores.',
    'Variable MISE_PYTHON_GITHUB_ATTESTATIONS=false en ambos entornos (fix build mise/Python 3.11.9).',
], 1):
    num(t)

d.add_page_break()

# ===== 2. ENTORNO =====
h1('2. Entorno y stack')
table(['Clave', 'Valor'], [
    ['Repositorio', 'github.com/Txolete/crm-backend'],
    ['Stack', 'FastAPI + SQLAlchemy 2.0 + PostgreSQL + Jinja2 + JS vanilla + Chart.js'],
    ['Deploy', 'Railway · main = PRO · develop = DEV'],
    ['Migraciones', 'Alembic (pre-deploy automático)'],
    ['IA', 'OpenAI Responses API · gpt-5.4 (general) · gpt-5.5 (comunicaciones)'],
    ['Roles', 'admin, sales, commercial, viewer'],
], widths=[1.6, 5.0])

# ===== 3. ESTADO REAL =====
h1('3. Estado real del proyecto (junio 2026)')
p('Tabla de sprints actualizada respecto a la v2.1. Cambios marcados en la columna Estado.')
table(['Sprint / Módulo', 'Descripción', 'Estado'], [
    ['Sprint 0-3', 'Deuda técnica, bugs, rol commercial, modal cierre, UX Kanban', 'Completado · PRO'],
    ['Sprint 4', 'Campos estratégicos, estado mental cliente, columna HOLD, motivos pérdida (migraciones aplicadas)', 'Completado · PRO (verificar UI)'],
    ['Sprint 5', 'Export Markdown, calendario tareas, .ics, campos ChatGPT', 'Parcial · PRO'],
    ['Bugs UX (sesión jun)', 'Fecha de tarea (prevalece manual) + calendario (preview hover + fix modal gris)', 'Completado · PRO'],
    ['Feedback in-app', 'Modelo UserFeedback, botón flotante global, panel admin', 'Completado · PRO'],
    ['Material comercial (N3+)', 'MaterialDocument con versionado (una activa, retirar/reactivar)', 'Completado · PRO'],
    ['Plantillas email comercial', 'EmailTemplate + EmailSent, 5 plantillas, secuencia 3 toques, firma usuario, CC/CCO, métricas', 'Completado · PRO'],
    ['N2 Hub comunicaciones BOMP', 'Ingesta Excel + IA + maqueta HTML correo', 'En DEV (esperando IT)'],
    ['N4 Ofertas (landing+PDF)', 'Generación de ofertas desde CRM', 'Funcionando · PRO'],
    ['N5 v1 Business plans', 'HTML standalone (a sustituir por N5 v2)', 'Funcionando · PRO'],
    ['Sprint 6', 'Email diario pipeline + KPI histórico', 'Pendiente'],
    ['Sprint 7', 'Oportunidades archivadas + dashboard actividad + docs', 'Pendiente'],
    ['N1 Onboarding', 'Microservicio multi-tenant + integración CRM', 'Pendiente (spec integrada)'],
    ['N5 v2 Business Plan Tool', 'Módulo integrado en CRM (sustituye N5 v1)', 'Pendiente (spec integrada)'],
], widths=[1.9, 3.9, 1.3])

# ===== 4. FUNCIONALIDADES NUEVAS DE LA SESION =====
h1('4. Funcionalidades completadas (junio 2026) — detalle')
p('Funcionalidades construidas fuera del scope original de la v2.1. Documentadas para que la otra '
  'licencia conozca los modelos existentes y NO los duplique.')

h2('4.1 Feedback in-app')
bullet('Modelo UserFeedback (app/models/feedback.py): id, user_id, message, view, url, user_agent, status (open|reviewed|done|discarded), created_at, reviewed_at, reviewed_by_user_id, admin_note.')
bullet('Botón flotante 💬 en todas las vistas (app/static/js/feedback_widget.js). Captura vista/URL/usuario automáticamente.')
bullet('Endpoints: POST /feedback, GET /feedback (admin todo, resto lo suyo), PATCH /feedback/{id}. Página admin /feedback/admin-page.')

h2('4.2 Material comercial')
bullet('Modelo MaterialDocument (app/models/material.py): archivo en bytea (Postgres, máx 25MB), name, name_slug, version, usage_note, status (active|retired).')
bullet('Regla: una sola versión activa por name_slug; al subir nueva, la anterior pasa a retired automáticamente.')
bullet('Endpoints en /materials: list, upload (admin), download, retire/restore, delete. Página /materials/page.')

h2('4.3 Plantillas de email comercial')
bullet('Modelos EmailTemplate + EmailSent (app/models/email_template.py). ESTA es la tabla de plantillas reutilizable para todo el CRM.')
bullet('5 plantillas seed: frío estándar, frío corporate/TIER1, y 3 toques de seguimiento (valor regulatorio / prueba social+deck / cierre limpio).')
bullet('Variables: nombre, empresa, senal_detectada (obligatoria configurable), firma_comercial, novedad_regulatoria, referencia, propuesta_hueco. Validación backend de obligatorias.')
bullet('Firma por usuario (users.email_signature). Multi-destinatarios + CC/CCO. Métricas enviados/respuestas.')
bullet('Secuencia de 3 toques: widget en ficha de oportunidad (reconstruido desde EmailSent por categoría de plantilla) + badge en Kanban. Marcado de respuesta manual.')
bullet('Endpoints: /email-templates (CRUD + render + seed-initial), /emails-sent (registro + respuesta), /opportunities/{id}/email-sequence.')

h2('4.4 N2 — Hub de comunicación de novedades BOMP (en DEV)')
bullet('Modelos (app/models/comunicacion.py): Publicacion (batch), Desarrollo (item en crudo del ERP), SalidaCanal (lo específico por canal, extensible), ComunicacionPrompt (versiones de prompt + hero_level 1-3 + calibración).')
bullet('Flujo: ingesta Excel de novedades → adaptar con IA (gpt-5.5, prompt versionable + feedback de calibración bien/meh/mal) → maqueta HTML email-safe (logo oficial, hero, CTA contacto, acento de color lateral por peso) → copiar con formato / descargar .eml (X-Unsent, editable en Outlook).')
bullet('Solo admin. Página /comunicaciones/page. Salida manual (sin SMTP en MVP).')
bullet('Pendiente IT (ver sección 7).')

d.add_page_break()

# ===== 5. N2 PENDIENTE IT =====
h1('5. N2 — Cierre del Hub de comunicaciones (depende de IT)')
h2('5.1 Export de novedades (recurrente, por release)')
p('IT amplía el export de novedades del ERP. Columnas:')
table(['Columna', 'Estado', 'Notas'], [
    ['ID', 'Nueva', 'Identificador único y estable del desarrollo en BOMP. Primera columna. Para dedupe y trazabilidad.'],
    ['Actualización', 'Existe', 'Título en crudo'],
    ['Tipo', 'Normalizar', 'EXACTAMENTE: Nueva funcionalidad / Mejora de funcionalidad existente / Adaptación regulatoria / Corrección de errores'],
    ['Fecha', 'Existe', ''],
    ['Observaciones', 'Existe', 'Texto técnico que adapta la IA'],
    ['Módulo', 'Existe', ''],
    ['Origen', 'Existe', 'Adm / Extranet / Ambas'],
    ['Proyecto', 'Existe', 'BOMP 1 / BOMP 2 / API / GAS...'],
    ['Norma', 'Nueva', 'Referencia legal + plazo (solo regulatorio)'],
    ['Mantenimiento', 'Nueva', 'Sí/No → agrupar sin titular'],
    ['Relacionado_con', 'Nueva', 'Referencia el ID de otro desarrollo (consolidar evolutivos)'],
], widths=[1.5, 1.0, 4.1])

h2('5.2 Export de clientes BOMP (una vez, para clonar fichas)')
p('Hoja "Clientes" (una fila por empresa):')
bullet('Nombre (oblig.), CIF/NIF (oblig. — clave anti-duplicados), Email empresa, Teléfono, Web, Dirección, Tipo de cliente, Provincia, Notas.')
p('Hoja "Contactos" (una fila por persona):')
bullet('CIF/NIF cliente (oblig. — enlace), Nombre, Apellidos, Rol, Email (clave para boletín), Teléfono.')
p('Campo recomendado (RGPD): "Acepta comunicaciones" Sí/No.')

h2('5.3 Trabajo CRM tras recibir los exports')
num('Ampliar parser de novedades (Norma, Mantenimiento, Relacionado_con) + guardar id_erp y resolver relacionados por ID.')
num('Crear importador de clientes/contactos (2 hojas, dedupe por CIF) → crea Account + Contact + canales email.')
num('Selector de destinatarios desde las fichas del CRM (sustituye pegado manual de emails).')
num('Mergear N2 a PRO y validar con una release real.')

d.add_page_break()

# ===== 6. N1 ONBOARDING =====
h1('6. N1 — Gestor de Onboarding (microservicio + integración CRM)')
p('Convierte la Onboarding Tool (Flask, hoy 1 cliente: Altano) en multi-tenant e integra su lanzamiento '
  'desde el CRM. Microservicio independiente; comunicación server-to-server por HTTP REST con token compartido.')
h2('6.1 Decisiones arquitectónicas')
for t in [
    'Microservicio independiente (Flask, repo/BD/deploy propios). Aislamiento de fallos.',
    'CRM → Onboarding: HTTP REST con header X-API-Token (ONBOARDING_API_TOKEN, mismo valor en ambos). El CRM hace pull del progreso, sin webhooks.',
    'Multi-tenancy path-based: /{slug} (ej. onboarding.../altano).',
    'Password único por cliente, generado al crear, mostrado una vez, hasheado bcrypt.',
    'Usuarios ASIC globales; asignaciones tarea↔responsable por cliente.',
    'Catálogo de tareas (DATA) hardcoded en frontend en esta iteración (a BD en v2).',
    'Lanzamiento desde oportunidad en Negociación Final o Ganada. Al pasar a Ganada, emergente sugiriendo lanzar (nunca automático).',
]:
    bullet(t)

h2('6.2 Modelo en CRM — tabla onboardings (FK como String/UUID)')
p('IMPORTANTE: la spec original usaba INT en los FK. En este CRM los IDs son VARCHAR/UUID. Corregido:')
table(['Campo', 'Tipo', 'Notas'], [
    ['id', 'String PK', 'UUID'],
    ['opportunity_id', 'String FK', '→ opportunities.id'],
    ['account_id', 'String FK', '→ accounts.id (denormalizado)'],
    ['external_slug', 'String UNIQUE', "'altano'"],
    ['external_url', 'String', 'URL completa al portal'],
    ['initial_password_shown', 'Integer (0/1)', '¿se mostró ya el password?'],
    ['status', 'String', 'active|completed|cancelled'],
    ['created_by_user_id', 'String FK', '→ users.id'],
    ['progress_total/done/pct', 'Integer', 'cache de progreso (job)'],
    ['last_synced_at', 'UTCDateTime', 'última sync'],
    ['created_at / completed_at', 'UTCDateTime', ''],
], widths=[1.9, 1.4, 3.3])

h2('6.3 Sub-sprints')
table(['Sub-sprint', 'Objetivo', 'Repo · Estim.'], [
    ['O1', 'Multi-tenant foundation: Alembic, tabla clients, routing path-based, migración de Altano', 'onboarding-tool · 2-3d'],
    ['O2', 'Usuarios ASIC, asignaciones tarea↔responsable, selector "¿quién eres?", filtro mis tareas', 'onboarding-tool · 2d'],
    ['O3', 'API externa con token (POST /onboardings, GET /progress, listado)', 'onboarding-tool · 1-2d'],
    ['CRM-N1', 'Modelo onboardings, botón "Lanzar onboarding" en oportunidad, modal password inicial', 'crm-backend · 1-2d'],
    ['CRM-N2', 'Vista /onboardings con KPI de progreso + job sync cada 15 min', 'crm-backend · 2d'],
], widths=[1.1, 4.2, 1.3])
p('Orden: O1 primero (cimientos). O2/O3 en paralelo. CRM-N1 requiere O3 en dev. CRM-N2 cierra el circuito.')

h2('6.4 Endpoints clave')
bullet('Onboarding (externa, X-API-Token): POST /api/external/onboardings · GET /api/external/onboardings/{slug}/progress · GET /api/external/onboardings · PATCH /api/external/onboardings/{slug}')
bullet('CRM: POST /opportunities/{id}/onboarding · GET /onboardings · GET /onboardings/{id}')
bullet('Trigger Ganada: PATCH /opportunities/{id} detecta paso a Ganada sin onboarding → respuesta con suggest_onboarding=true (el frontend abre el emergente; nunca crea solo).')

h2('6.5 Variables de entorno nuevas (CRM)')
bullet('ONBOARDING_BASE_URL=https://onboarding-tool.up.railway.app')
bullet('ONBOARDING_API_TOKEN=mismo valor que en Onboarding Tool')

d.add_page_break()

# ===== 7. N5 v2 BUSINESS PLAN =====
h1('7. N5 v2 — Business Plan Tool integrado')
p('Sustituye el N5 v1 (HTML standalone tipo Valfortec enviado a mano) por un módulo integrado en el CRM, '
  'con persistencia y contexto. "Hacerlo como Dios manda": el simulador vive dentro del CRM, vinculado a '
  'cuenta/oportunidad, con múltiples simulaciones, y genera un HTML blindado para el cliente. Patrón N4.')
h2('7.1 Alcance')
bullet('Modelo financiero simple: 3 años, KPIs Facturación / Margen bruto / EBITDA, 3 escenarios (conservador/base/ambicioso). Idéntico al Valfortec actual.')
bullet('NO entra: modelo a 5 años (P&L completa, garantías OMIE/MEFF) — herramienta aparte post-contratación. NI formulario web público (el comercial mete los datos a mano). NI microservicio. NI envío SMTP automático.')
h2('7.2 Modelo de datos — tabla business_plans')
p('Una sola tabla nueva. Las plantillas de correo REUTILIZAN email_templates (decisión: NO crear cfg_email_templates).')
table(['Campo', 'Tipo', 'Notas'], [
    ['id', 'String PK', 'UUID'],
    ['account_id', 'String FK', 'obligatorio (denormalizado desde opp si nace de oportunidad)'],
    ['opportunity_id', 'String FK', 'opcional (NULL si nace de cuenta)'],
    ['title', 'String', 'autogenerado: "BP {cuenta} - {escenario}"'],
    ['scenario', 'String', 'conservador|base|ambicioso|NULL (personalizado)'],
    ['inputs_json', 'Text/JSON', 'variables del simulador (flexible, sin migración en v2)'],
    ['kpis_cache', 'Text/JSON', 'KPIs cacheados (recalcular si cambia fórmula)'],
    ['status', 'String', 'draft|sent|archived'],
    ['owner_user_id', 'String FK', 'default = created_by; commercial solo ve los suyos'],
    ['created_by_user_id', 'String FK', ''],
    ['sent_at / archived_at / created_at / updated_at', 'UTCDateTime', ''],
], widths=[2.4, 1.3, 2.9])

h2('7.3 Plantillas de correo (REUTILIZAR email_templates)')
bullet('NO crear cfg_email_templates. Usar la tabla existente email_templates con plantillas cuyo category empiece por "bp_" (p. ej. bp_initial_request).')
bullet('Seed de la plantilla bp_initial_request (solicitud de datos al potencial, basada en el correo a Valfortec) insertada como EmailTemplate.')
bullet('Placeholders ya soportados por el motor de plantillas: nombre/empresa/firma_comercial + añadir los del BP (account_name, account_region, opportunity_name, commercial_email, commercial_phone, today_date) al render.')

h2('7.4 UI')
bullet('Vista /business-plans: listado con filtros (status, cuenta, comercial, búsqueda) y filtro de rol (commercial solo los suyos).')
bullet('Vista /business-plans/{id}: simulador 2 columnas (inputs / outputs con Chart.js), replica del Valfortec. Cálculo en cliente en tiempo real; guardar persiste inputs_json + KPIs.')
bullet('Secciones "Business Plans" en ficha de cuenta y de oportunidad (botón "+ Lanzar BP").')
bullet('Modal "Generar correo": render de plantilla (email_templates), copiar al portapapeles / abrir en cliente de correo (mailto:) / marcar como enviado.')
h2('7.5 HTML blindado para el cliente')
bullet('GET /business-plans/{id}/export.html: fichero autocontenido (CSS+JS embebidos, solo Chart.js por CDN), inputs pre-cargados con value="...", lógica de cálculo fija. Branding ASIC. Sin IDs internos.')
bullet('Nombre: BP_{slug(cuenta)}_{slug(title)}_{YYYYMMDD}.html. Cada descarga registra audit_log (export_html).')
h2('7.6 Sprints N5v2')
bullet('N5v2-1: backend + modelo business_plans + seed plantilla en email_templates + endpoints.')
bullet('N5v2-2: UI simulador (listado + detalle + integración en fichas).')
bullet('N5v2-3: export HTML blindado + modal correo + render de plantilla.')

d.add_page_break()

# ===== 8. CONFLICTOS RESUELTOS =====
h1('8. Registro de decisiones (conflictos entre specs resueltos)')
table(['Tema', 'Decisión para v3'], [
    ['Plantillas de correo del Business Plan', 'Reutilizar email_templates (categoría bp_*). NO crear cfg_email_templates.'],
    ['N5 v1 vs N5 v2', 'Lo que hay en PRO es el HTML standalone. N5 v2 es el sprint nuevo que lo sustituye como módulo integrado.'],
    ['FK del modelo onboardings (y business_plans)', 'Todos los IDs del CRM son VARCHAR/UUID. Los FK nuevos se definen como String, no INT.'],
    ['Múltiples "sistemas de correo"', 'Unificar en email_templates como tabla de plantillas reutilizable. comunicacion_prompts es distinto (prompts de IA del hub, no plantillas de correo).'],
], widths=[2.3, 4.3])

# ===== 9. MIGRACIONES =====
h1('9. Migraciones Alembic')
p('Aplicar siempre en dev primero, validar, y después en pro. Comando: alembic upgrade head')
h3('Ya aplicadas (junio 2026, no estaban en la tabla de la v2.1)')
for t in ['add_cfg_opportunity_types_lost_reasons_mental_states', 'add_opportunity_new_fields', 'add_hold_stage', 'add_opportunity_outcomes', 'add_ai_chat_history_external_notes', 'add_created_by_user_id_to_tasks', 'add_user_feedback', 'add_material_documents', 'add_email_templates', 'user_signature + email cc/bcc + contact_roles', 'add_comunicaciones_hub', 'add_comunicacion_prompts']:
    bullet(t)
h3('Pendientes (sprints futuros)')
for t in ['create_onboardings (N1 CRM-N1)', 'create_business_plans (N5v2-1)', 'add_kpi_snapshots (Sprint 6)']:
    bullet(t)

# ===== 10. ORDEN =====
h1('10. Orden recomendado de ejecución')
table(['#', 'Sprint', 'Justificación'], [
    ['1', 'N2 cierre', 'Exports IT + importador clientes + selector destinatarios → mergear a PRO. Casi listo.'],
    ['2', 'N1 Onboarding (O1→O2/O3→CRM-N1→CRM-N2)', 'Microservicio multi-tenant + integración. Alto valor operativo.'],
    ['3', 'N5 v2 Business Plan', 'Módulo integrado reutilizando email_templates.'],
    ['4', 'Sprint 6', 'Email diario pipeline + KPI histórico. Infra casi lista.'],
    ['5', 'Sprint 7', 'Oportunidades archivadas + dashboard actividad + docs.'],
    ['6', 'Fase 2 comunicaciones/email', 'LinkedIn, SMTP directo, secuencias automáticas, push ERP API.'],
], widths=[0.4, 2.4, 3.8])

d.add_paragraph()
end = d.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.add_run('— Fin de la especificación v3.0 —').italic = True

out = 'CRM_ASICXXI_Spec_ClaudeCode_v3.0.docx'
d.save(out)
print('GENERADO:', out)
print('parrafos:', len(d.paragraphs), '| tablas:', len(d.tables))
